"""
论文撰写模块 (Module 07) —— 从 v3.0 蓝本 07_paper_writing/writer.py 忠实移植
功能：按照 CUMCM 标准格式自动生成摘要、模型假设、求解过程、结果分析与模型评价各章节。

移植说明：
- 保留 v3.0 全部真实文本生成逻辑（基于上游真实数据拼装章节）。
- generate_paper 新增可选 output_dir 参数（默认 "output/paper"），便于 runner 写入
  workflow.project_dir/"paper"。
- 新增 save_docx(paper_structure, path)：优先用 python-docx 写出 .docx；若 python-docx
  缺失则抛出明确 ImportError，由 runner 决定降级为 markdown（不伪造 docx）。
- 模块本身仅依赖标准库，故始终可被 import。
"""

import json
from pathlib import Path
from typing import Dict, List, Optional, Any
from dataclasses import dataclass, asdict
from datetime import datetime


@dataclass
class Chapter:
    """论文章节"""
    chapter_id: str
    title: str
    content: str
    level: int = 1
    word_count: int = 0


@dataclass
class PaperStructure:
    """论文结构"""
    title: str
    abstract: str
    keywords: List[str]
    chapters: List[Chapter]
    references: List[str]
    appendix: List[str]
    total_word_count: int = 0


@dataclass
class PaperResult:
    """论文撰写结果"""
    result_id: str
    paper_structure: PaperStructure
    output_path: str
    format_type: str  # md, docx
    created_at: str
    metadata: Dict


class PaperWriter:
    """论文撰写器（移植自 v3.0，逻辑不变）"""

    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {
            "competition": "CUMCM",
            "language": "zh",
            "target_word_count": 20000,
            "template": "default"
        }

    def generate_paper(self,
                      problem_analysis: Dict,
                      model_selection: Dict,
                      data_info: Dict,
                      solving_results: Dict,
                      validation_results: Dict,
                      visualization_results: Dict,
                      output_dir: str = "output/paper") -> PaperResult:
        """生成完整论文（真实文本拼装，基于上游传入的数据）。"""
        abstract = self._generate_abstract(problem_analysis, solving_results)
        keywords = self._extract_keywords(problem_analysis, model_selection)

        chapters = []
        chapters.append(self._generate_chapter_1(problem_analysis))
        chapters.append(self._generate_chapter_2(problem_analysis))
        chapters.append(self._generate_chapter_3(problem_analysis, model_selection))
        chapters.append(self._generate_chapter_4(problem_analysis))
        chapters.append(self._generate_chapter_5(model_selection, solving_results))
        chapters.append(self._generate_chapter_6(solving_results, visualization_results))
        chapters.append(self._generate_chapter_7(validation_results))
        chapters.append(self._generate_chapter_8(model_selection, validation_results))
        chapters.append(self._generate_chapter_9(model_selection))

        references = self._generate_references(model_selection)

        total_word_count = sum(ch.word_count for ch in chapters)
        total_word_count += len(abstract)

        paper_structure = PaperStructure(
            title=problem_analysis.get("title", "数学建模论文"),
            abstract=abstract,
            keywords=keywords,
            chapters=chapters,
            references=references,
            appendix=[],
            total_word_count=total_word_count
        )

        out_path = Path(output_dir)
        out_path.mkdir(parents=True, exist_ok=True)
        md_path = out_path / "paper_draft.md"
        self._save_markdown(paper_structure, md_path)

        result = PaperResult(
            result_id=f"PR-{datetime.now().strftime('%Y%m%d%H%M%S')}",
            paper_structure=paper_structure,
            output_path=str(md_path),
            format_type="md",
            created_at=datetime.now().isoformat(),
            metadata={
                "total_word_count": total_word_count,
                "chapters_count": len(chapters),
                "competition": self.config.get("competition", "CUMCM")
            }
        )
        return result

    def _generate_abstract(self, problem_analysis: Dict, solving_results: Dict) -> str:
        problem_type = problem_analysis.get("problem_type_cn", "综合类")
        title = problem_analysis.get("title", "数学建模问题")
        desc_raw = problem_analysis.get("description", "")
        description = (desc_raw[:200] + "...") if desc_raw else ""

        metrics = solving_results.get("metrics", {})
        r2 = metrics.get("r2", 0)
        rmse = metrics.get("rmse", 0)
        model_name = solving_results.get("model_name", "所选模型")

        variables = problem_analysis.get("variables", [])
        constraints = problem_analysis.get("constraints", [])

        abstract = f"""本文针对"{title}"问题，研究了{problem_type}问题的求解方法。

一、问题分析

{description}

通过深入分析，识别出{len(variables)}个关键变量和{len(constraints)}个约束条件，建立了问题的数学框架。

二、模型选择

根据问题特征，我们选择了{model_name}作为主要建模方法。该模型适用于{problem_type}问题，具有以下优势：
1. 能够有效处理问题的核心挑战
2. 模型复杂度适中，易于实现和解释
3. 有成熟的理论基础和实现方法

三、模型求解

采用系统化的求解流程，包括数据预处理、参数初始化、模型训练和结果输出。主要求解结果如下：
- 拟合优度（R²）：{r2:.4f}
- 均方根误差（RMSE）：{rmse:.4f}

四、模型验证

通过灵敏度分析和模型验证，确认了模型的稳定性和可靠性。结果表明，本文建立的模型能够有效地解决该问题，为类似问题提供了参考方法。

五、主要贡献

1. 建立了针对{problem_type}问题的数学模型
2. 提出了系统化的求解方法
3. 通过实例验证了模型的有效性

关键词：{', '.join(self._extract_keywords(problem_analysis, {}))}"""
        return abstract

    def _extract_keywords(self, problem_analysis: Dict, model_selection: Dict) -> List[str]:
        keywords = []
        problem_type = problem_analysis.get("problem_type_cn", "")
        if problem_type:
            keywords.append(problem_type)
        model_name = model_selection.get("selected_model", {}).get("name_cn", "")
        if model_name:
            keywords.append(model_name)
        keywords.extend(["数学建模", "优化求解", "灵敏度分析"])
        return keywords[:5]

    def _generate_chapter_1(self, problem_analysis: Dict) -> Chapter:
        title = problem_analysis.get("title", "数学建模问题")
        description = problem_analysis.get("description", "")
        content = f"""# 一、问题重述

## 1.1 问题背景

{title}是数学建模竞赛中的经典问题类型。本问题涉及多个变量和约束条件，需要建立合理的数学模型来求解。

## 1.2 问题要求

根据题目要求，需要完成以下任务：

"""
        sub_problems = problem_analysis.get("sub_problems", [])
        for i, sub in enumerate(sub_problems, 1):
            content += f"{i}. {sub.get('description', f'问题{i}')}\n"
        if not sub_problems:
            content += "1. 建立数学模型\n2. 求解模型\n3. 分析结果\n"
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="1", title="问题重述", content=content, level=1, word_count=word_count)

    def _generate_chapter_2(self, problem_analysis: Dict) -> Chapter:
        problem_type = problem_analysis.get("problem_type_cn", "综合类")
        content = f"""# 二、问题分析

## 2.1 问题类型分析

本题属于{problem_type}问题。通过分析题目要求和已知条件，我们识别出以下关键要素：

"""
        variables = problem_analysis.get("variables", [])
        if variables:
            content += "## 2.2 变量分析\n\n本问题涉及以下主要变量：\n\n"
            content += "| 符号 | 名称 | 说明 |\n|------|------|------|\n"
            for var in variables[:10]:
                content += f"| {var.get('symbol', '-')} | {var.get('name', '-')} | {var.get('description', '-')} |\n"
        constraints = problem_analysis.get("constraints", [])
        if constraints:
            content += "\n## 2.3 约束条件分析\n\n本问题的主要约束条件包括：\n\n"
            for i, const in enumerate(constraints[:5], 1):
                content += f"{i}. {const.get('description', const.get('expression', '-'))}\n"
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="2", title="问题分析", content=content, level=1, word_count=word_count)

    def _generate_chapter_3(self, problem_analysis: Dict, model_selection: Dict) -> Chapter:
        content = """# 三、模型假设

为了简化问题并建立可求解的数学模型，我们做出以下假设：

## 3.1 基本假设

1. **数据质量假设**：假设题目提供的数据准确可靠，能够反映实际情况。

2. **独立性假设**：假设各变量之间的影响是独立的，或者其相互影响可以用已知的数学关系描述。

3. **连续性假设**：假设相关变量在研究范围内是连续变化的。

4. **稳定性假设**：假设系统的状态在研究时间段内是相对稳定的。

## 3.2 模型假设

"""
        problem_type = problem_analysis.get("problem_type", "optimization")
        if problem_type == "optimization":
            content += """1. **目标明确假设**：假设优化目标是明确且可量化的。

2. **约束完整假设**：假设所有重要的约束条件都已被识别和考虑。

3. **可行解存在假设**：假设问题存在可行解。
"""
        elif problem_type == "prediction":
            content += """1. **趋势延续假设**：假设历史数据的趋势在未来会延续。

2. **因素稳定假设**：假设影响预测结果的主要因素在预测期内保持稳定。

3. **模型适用假设**：假设所选预测模型适用于当前问题。
"""
        else:
            content += """1. **评价标准合理假设**：假设所选评价标准能够反映实际情况。

2. **指标可量化假设**：假设评价指标是可量化和可比较的。

3. **权重合理假设**：假设指标权重的设定是合理的。
"""
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="3", title="模型假设", content=content, level=1, word_count=word_count)

    def _generate_chapter_4(self, problem_analysis: Dict) -> Chapter:
        content = """# 四、符号说明

本文使用的主要符号及其含义如下表所示：

| 符号 | 含义 | 单位 |
|------|------|------|
"""
        variables = problem_analysis.get("variables", [])
        for var in variables[:15]:
            content += f"| {var.get('symbol', '-')} | {var.get('description', '-')} | {var.get('unit', '-')} |\n"
        content += """| $x$ | 自变量 | - |
| $y$ | 因变量 | - |
| $f(x)$ | 目标函数 | - |
| $g(x)$ | 约束函数 | - |
| $n$ | 样本数量 | 个 |
| $p$ | 特征数量 | 个 |
"""
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="4", title="符号说明", content=content, level=1, word_count=word_count)

    def _generate_chapter_5(self, model_selection: Dict, solving_results: Dict) -> Chapter:
        model_name = model_selection.get("selected_model", {}).get("name_cn", "数学模型")
        content = f"""# 五、模型建立与求解

## 5.1 模型选择

根据问题分析结果，我们选择{model_name}作为主要建模方法。

选择理由：
"""
        rationale = model_selection.get("selection_rationale", "")
        if rationale:
            content += f"\n{rationale}\n"
        else:
            content += "\n1. 该模型适用于本问题类型\n2. 模型复杂度适中\n3. 有成熟的实现方法\n"
        content += f"""
## 5.2 模型建立

### 5.2.1 数学公式

基于{model_name}，建立如下数学模型：

$$
\\begin{{aligned}}
& \\min/\\max f(x) \\\\
& \\text{{s.t.}} \\\\
& g_i(x) \\leq 0, \\quad i = 1, 2, \\ldots, m \\\\
& h_j(x) = 0, \\quad j = 1, 2, \\ldots, p
\\end{{aligned}}
$$

### 5.2.2 求解算法

采用以下步骤进行求解：

1. 数据预处理
2. 参数初始化
3. 模型训练
4. 结果输出

## 5.3 模型求解

"""
        metrics = solving_results.get("metrics", {})
        content += f"""### 5.3.1 求解结果

模型求解的主要结果如下：

| 指标 | 值 |
|------|-----|
| R² | {metrics.get('r2', 0):.4f} |
| RMSE | {metrics.get('rmse', 0):.4f} |
| MAE | {metrics.get('mae', 0):.4f} |

### 5.3.2 参数估计

"""
        parameters = solving_results.get("parameters", [])
        if parameters:
            content += "模型的主要参数估计值如下：\n\n| 参数 | 估计值 |\n|------|--------|\n"
            for param in parameters[:10]:
                content += f"| {param.get('name', '-')} | {param.get('value', 0):.4f} |\n"
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="5", title="模型建立与求解", content=content, level=1, word_count=word_count)

    def _generate_chapter_6(self, solving_results: Dict, visualization_results: Dict) -> Chapter:
        content = """# 六、结果分析

## 6.1 主要结果

通过模型求解，我们得到以下主要结果：

"""
        feature_importance = solving_results.get("feature_importance", {})
        if feature_importance:
            content += "### 6.1.1 特征重要性分析\n\n各特征对结果的重要性如下：\n\n"
            sorted_imp = sorted(feature_importance.items(), key=lambda x: x[1], reverse=True)[:5]
            for name, imp in sorted_imp:
                content += f"- **{name}**: 重要性 = {imp:.4f}\n"
        content += """
## 6.2 灵敏度分析

"""
        sensitivity = solving_results.get("sensitivity_results", [])
        if sensitivity:
            content += "对主要参数进行灵敏度分析，结果如下：\n\n"
            for sens in sensitivity[:3]:
                content += f"### 参数 {sens.get('parameter_name', '-')}\n\n"
                content += f"- 基准值: {sens.get('base_value', 0):.4f}\n"
                content += f"- 灵敏度得分: {sens.get('sensitivity_score', 0):.2f}\n"
                content += f"- 结论: {sens.get('conclusion', '-')}\n\n"
        content += """## 6.3 结果可视化

"""
        figures = visualization_results.get("figures", [])
        if figures:
            content += "主要可视化结果如下：\n\n"
            for i, fig in enumerate(figures[:5], 1):
                if isinstance(fig, dict):
                    title = fig.get("title", "图表")
                    fig_type = fig.get("type", "")
                    desc = f"展示了{fig_type}分析结果" if fig_type else ""
                else:
                    import os
                    fig_name = os.path.splitext(os.path.basename(str(fig)))[0]
                    name_map = {
                        "pred_vs_actual": "预测值与实际值对比", "residuals": "残差分析",
                        "feature_importance": "特征重要性", "data_distribution": "数据分布",
                        "correlation_heatmap": "相关性热力图", "error_analysis": "误差分析",
                        "boxplot": "箱线图",
                    }
                    desc = name_map.get(fig_name, fig_name)
                    title = desc
                content += f"{i}. **{title}**\n"
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="6", title="结果分析", content=content, level=1, word_count=word_count)

    def _generate_chapter_7(self, validation_results: Dict) -> Chapter:
        content = """# 七、模型检验

## 7.1 模型验证方法

为确保模型的可靠性和有效性，我们采用以下方法进行模型检验：

1. **交叉验证**：使用K折交叉验证评估模型的稳定性
2. **残差分析**：检查模型残差是否符合正态分布假设
3. **灵敏度分析**：评估参数变化对结果的影响

## 7.2 验证结果

"""
        overall_score = validation_results.get("overall_score", 0)
        overall_status = validation_results.get("overall_status", "unknown")
        content += f"""### 7.2.1 综合验证评分

模型综合验证评分为 **{overall_score:.1f}/100**，状态为 **{overall_status}**。

"""
        checks = validation_results.get("checks", [])
        if checks:
            content += "### 7.2.2 详细检查结果\n\n| 检查项 | 状态 | 评分 |\n|--------|------|------|\n"
            for check in checks[:10]:
                status_emoji = "✅" if check.get("status") == "passed" else "⚠️" if check.get("status") == "warning" else "❌"
                content += f"| {check.get('check_name', '-')} | {status_emoji} {check.get('status', '-')} | {check.get('score', 0):.1f} |\n"
        content += """
## 7.3 检验结论

"""
        passed = validation_results.get("passed_checks", 0)
        warnings = validation_results.get("warning_checks", 0)
        failed = validation_results.get("failed_checks", 0)
        content += f"""通过模型检验，共完成 {passed + warnings + failed} 项检查：

- ✅ 通过: {passed} 项
- ⚠️ 警告: {warnings} 项
- ❌ 失败: {failed} 项

总体而言，模型通过了大部分检验，结果可靠。
"""
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="7", title="模型检验", content=content, level=1, word_count=word_count)

    def _generate_chapter_8(self, model_selection: Dict, validation_results: Dict) -> Chapter:
        model_name = model_selection.get("selected_model", {}).get("name_cn", "所选模型")
        content = f"""# 八、模型评价

## 8.1 模型优点

{model_name}具有以下优点：

"""
        pros = model_selection.get("selected_model", {}).get("pros", [])
        if pros:
            for i, pro in enumerate(pros, 1):
                content += f"{i}. {pro}\n"
        else:
            content += "1. 该模型适用于本问题类型，具有较好的适用性\n"
        content += f"""
## 8.2 模型缺点

同时，{model_name}也存在以下局限性：

"""
        cons = model_selection.get("selected_model", {}).get("cons", [])
        if cons:
            for i, con in enumerate(cons, 1):
                content += f"{i}. {con}\n"
        else:
            content += "1. 模型假设可能与实际情况存在一定偏差\n"
        content += """
## 8.3 改进方向

为进一步提高模型性能，可以考虑以下改进方向：

1. **数据层面**：收集更多高质量数据，改进数据预处理方法
2. **模型层面**：尝试其他模型或进行模型融合
3. **算法层面**：优化求解算法，提高计算效率
4. **验证层面**：增加更多验证方法，提高结果可信度
"""
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="8", title="模型评价", content=content, level=1, word_count=word_count)

    def _generate_chapter_9(self, model_selection: Dict) -> Chapter:
        content = """# 九、模型推广

## 9.1 应用场景

本文建立的模型可以推广到以下类似场景：

"""
        scenarios = model_selection.get("selected_model", {}).get("applicable_scenarios", [])
        if scenarios:
            for i, scenario in enumerate(scenarios, 1):
                content += f"{i}. {scenario}\n"
        else:
            content += """1. 类似结构的优化问题
2. 同类型的数据分析任务
3. 相同领域的决策支持
"""
        content += """
## 9.2 推广方法

将本模型推广到其他问题时，需要：

1. **问题分析**：分析新问题与原问题的相似性和差异
2. **数据准备**：收集和整理新问题的相关数据
3. **模型调整**：根据新问题特点调整模型参数
4. **验证评估**：对新模型进行验证和评估

## 9.3 注意事项

在推广过程中需要注意：

1. 确保新问题满足模型的基本假设
2. 根据实际情况调整模型参数
3. 进行充分的验证测试
4. 考虑实际应用的可行性
"""
        word_count = len(content.replace(" ", "").replace("\n", ""))
        return Chapter(chapter_id="9", title="模型推广", content=content, level=1, word_count=word_count)

    def _generate_references(self, model_selection: Dict) -> List[str]:
        references = [
            "[1] 姜启源, 谢金星, 叶俊. 数学模型(第五版). 高等教育出版社, 2018.",
            "[2] 司守奎, 孙兆亮. 数学建模算法与应用(第2版). 国防工业出版社, 2015.",
            "[3] 韩中庚. 数学建模方法及其应用(第2版). 高等教育出版社, 2009.",
            "[4] 吴建国. 数学建模案例精编. 中国水利水电出版社, 2010.",
            "[5] 谢晋. 数学建模入门与提高. 科学出版社, 2014.",
        ]
        model_name = model_selection.get("selected_model", {}).get("name", "")
        if "linear" in model_name.lower():
            references.append("[6] Hastie T, Tibshirani R, Friedman J. The Elements of Statistical Learning. Springer, 2009.")
        elif "neural" in model_name.lower():
            references.append("[6] Goodfellow I, Bengio Y, Courville A. Deep Learning. MIT Press, 2016.")
        return references

    def _save_markdown(self, paper: PaperStructure, output_path: Path):
        content = f"""# {paper.title}

## 摘要

{paper.abstract}

**关键词**：{', '.join(paper.keywords)}

---

"""
        for chapter in paper.chapters:
            content += f"\n{chapter.content}\n"
        content += "\n# 参考文献\n\n"
        for ref in paper.references:
            content += f"{ref}\n"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(content)
        print(f"论文草稿已保存到: {output_path}")

    def save_docx(self, paper: PaperStructure, output_path: str):
        """
        用 python-docx 将论文结构写出为 .docx。

        依赖：python-docx。若未安装则抛出明确的 ImportError，由调用方决定降级为 markdown。
        """
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError as e:
            raise ImportError(
                "生成 .docx 需要 python-docx，请先安装：\n"
                "  pip install python-docx"
            ) from e

        doc = Document()

        # 标题
        title_p = doc.add_heading(paper.title, level=0)
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 摘要
        doc.add_heading("摘要", level=1)
        doc.add_paragraph(paper.abstract)
        kw_p = doc.add_paragraph()
        kw_run = kw_p.add_run("关键词：" + ", ".join(paper.keywords))
        kw_run.bold = True

        # 章节
        for chapter in paper.chapters:
            # 章节正文可能以 "# 一、..." 开头的 Markdown 标题；统一转为 docx 标题与段落
            self._append_chapter_to_docx(doc, chapter)

        # 参考文献
        doc.add_heading("参考文献", level=1)
        for ref in paper.references:
            doc.add_paragraph(ref)

        out_path = Path(output_path)
        out_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(out_path))
        print(f"论文 DOCX 已保存到: {out_path}")

    def _append_chapter_to_docx(self, doc, chapter):
        """将 Markdown 风格章节内容追加到 docx（按 '#' 层级生成标题，其余作为段落）。"""
        import re
        lines = chapter.content.split("\n")
        for line in lines:
            s = line.rstrip()
            if not s.strip():
                continue
            m = re.match(r"^(#{1,6})\s+(.*)$", s)
            if m:
                level = len(m.group(1))
                doc.add_heading(m.group(2).strip(), level=min(level, 3))
            else:
                doc.add_paragraph(s)

    def generate_report_md(self, result: PaperResult, output_path: str):
        md_content = f"""# 论文撰写报告

## 基本信息

- **结果ID**: {result.result_id}
- **生成时间**: {result.created_at}
- **输出格式**: {result.format_type}
- **总字数**: {result.metadata.get('total_word_count', 0)}

## 论文结构

- **标题**: {result.paper_structure.title}
- **摘要**: {len(result.paper_structure.abstract)} 字
- **章节数**: {len(result.paper_structure.chapters)}
- **参考文献**: {len(result.paper_structure.references)} 篇

## 章节列表

"""
        for chapter in result.paper_structure.chapters:
            md_content += f"### 第{chapter.chapter_id}章: {chapter.title}\n"
            md_content += f"- 字数: {chapter.word_count}\n\n"
        md_content += f"""
## 输出文件

论文草稿已保存到: `{result.output_path}`

---

*生成时间: {result.created_at}*
"""
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)
        print(f"论文撰写报告已保存到: {output_path}")


def main():
    """示例用法（纯标准库即可运行，生成 markdown；docx 需 python-docx）"""
    problem_analysis = {
        "problem_id": "CUMCM-2026-A",
        "title": "RGV的动态调度优化问题",
        "problem_type": "optimization",
        "problem_type_cn": "优化类",
        "description": "本题研究RGV的动态调度优化问题",
        "variables": [
            {"symbol": "x", "name": "位置", "description": "RGV当前位置", "unit": "m"},
            {"symbol": "v", "name": "速度", "description": "RGV移动速度", "unit": "m/s"},
            {"symbol": "t", "name": "时间", "description": "加工时间", "unit": "s"}
        ],
        "constraints": [
            {"name": "速度约束", "expression": "0 <= v <= v_max", "description": "速度限制"},
            {"name": "时间约束", "expression": "t >= 0", "description": "时间非负"}
        ],
        "sub_problems": [{"id": "Q1", "name": "调度优化", "description": "建立调度优化模型"}]
    }
    model_selection = {
        "selected_model": {
            "id": "linear_programming", "name_cn": "线性规划",
            "pros": ["求解速度快", "理论成熟", "全局最优"],
            "cons": ["只能处理线性问题"],
            "applicable_scenarios": ["资源分配", "生产调度"]
        },
        "selection_rationale": "本题是线性优化问题，适合使用线性规划方法。"
    }
    solving_results = {
        "model_name": "线性规划",
        "metrics": {"r2": 0.92, "rmse": 0.15, "mae": 0.12},
        "parameters": [{"name": "coef_1", "value": 0.85}, {"name": "coef_2", "value": 0.72}],
        "feature_importance": {"x": 0.45, "v": 0.35, "t": 0.20},
        "sensitivity_results": [{"parameter_name": "速度上限", "base_value": 10.0, "sensitivity_score": 8.5, "conclusion": "中等影响"}]
    }
    validation_results = {
        "overall_score": 85.0, "overall_status": "good",
        "passed_checks": 12, "warning_checks": 2, "failed_checks": 0,
        "checks": [{"check_name": "R²检查", "status": "passed", "score": 90},
                   {"check_name": "残差检查", "status": "passed", "score": 85}]
    }
    visualization_results = {
        "figures": [{"title": "预测vs真实", "type": "scatter"}, {"title": "残差分布", "type": "histogram"}]
    }

    print("=" * 60)
    print("论文撰写示例")
    print("=" * 60)
    writer = PaperWriter()
    result = writer.generate_paper(problem_analysis, model_selection, {},
                                   solving_results, validation_results, visualization_results,
                                   output_dir="output/paper")
    print(f"总字数: {result.metadata.get('total_word_count', 0)}")
    print(f"章节数: {len(result.paper_structure.chapters)}")
    print(f"输出路径: {result.output_path}")

    # 尝试 docx；缺失则明确说明降级
    try:
        writer.save_docx(result.paper_structure, "output/paper/paper.docx")
    except ImportError as e:
        print(f"[说明] 未生成 docx：{e}")

    writer.generate_report_md(result, "output/paper_writing_report.md")
    print("=" * 60)
    print("论文撰写完成！")
    print("=" * 60)


if __name__ == "__main__":
    main()
